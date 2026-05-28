# src/loaders/docx_loader.py

"""DOCX loader via ``python-docx``."""

from __future__ import annotations

from pathlib import Path

from src.loaders.base import BaseLoader, LoadedDocument, LoadedPage, LoaderError


class DOCXLoader(BaseLoader):
    """
    Load .docx as a single page of concatenated paragraphs + tables.

    Page-level provenance is unavailable from .docx (Word doesn't store
    print-page boundaries reliably), so we emit a single ``LoadedPage``
    with ``page_number=None``. Headings/tables/paragraphs are flattened
    in document order.
    """

    name = "docx"
    extensions = (".docx",)
    mime_types = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",)

    def load(self, path: Path) -> LoadedDocument:
        path = self._resolve(path)

        try:
            import docx  # python-docx
        except ImportError as exc:  # pragma: no cover - dependency missing
            raise LoaderError(f"python-docx is required to load {path}") from exc

        try:
            doc = docx.Document(str(path))
        except Exception as exc:
            raise LoaderError(f"Could not open docx {path}: {exc}") from exc

        # Walk the document body in order. python-docx exposes paragraphs
        # and tables as separate collections; iterating the underlying
        # XML element preserves the in-document order.
        parts = []
        try:
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
        except Exception as exc:
            raise LoaderError(f"Error walking docx {path}: {exc}") from exc

        text = "\n".join(parts)

        doc_metadata = {}
        try:
            core = doc.core_properties
            for attr in ("author", "title", "subject", "created", "modified"):
                value = getattr(core, attr, None)
                if value:
                    doc_metadata[attr] = str(value)
        except Exception:  # pragma: no cover - core_props may be absent
            pass

        return LoadedDocument(
            source_path=str(path),
            document_name=path.name,
            mime_type=self._mime_for(path) or self.mime_types[0],
            pages=[LoadedPage(text=text, page_number=None,
                              metadata=self._base_page_metadata(path))],
            metadata=doc_metadata,
        )
